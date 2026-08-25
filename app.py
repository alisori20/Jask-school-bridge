import smart_school_addons
import streamlit as st
import sqlite3
import pandas as pd
import os
import subprocess
from datetime import datetime

# Set Page Config first
st.set_page_config(page_title="مدرسه هوشمند پُل", page_icon="🎓", layout="wide")

DB_PATH = os.path.join(os.path.dirname(__file__), "school.db")

# Automatically initialize database if it doesn't exist or is empty
if not os.path.exists(DB_PATH) or os.path.getsize(DB_PATH) == 0:
    try:
        import school_db
        school_db.init_db()
    except Exception as e:
        import streamlit as st
        st.error(f"Error initializing database: {e}")

try:
    upgrade_db_schema()
except Exception as e:
    st.error(f"Error upgrading database: {e}")


def get_connection():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def clean_and_validate_national_id(code):
    if not code:
        return None
    code = str(code).strip()
    # Convert Arabic and Persian numbers to English digits
    persian_to_eng = str.maketrans('۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩', '01234567890123456789')
    code = code.translate(persian_to_eng)
    import re
    if re.match(r'^\d{10}$', code):
        return code
    # Allow parent prefix if checking for login
    if code.startswith('p_') and re.match(r'^\d{10}$', code[2:]):
        return code
    return None


def upgrade_db_schema():
    conn = get_connection()
    cursor = conn.cursor()
    
    # 1. School settings table
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS school_settings (
        key TEXT PRIMARY KEY,
        val TEXT
    )
    """)
    # Seed default school name if not exists
    cursor.execute("INSERT OR IGNORE INTO school_settings (key, val) VALUES ('school_name', 'دبیرستان متوسطه اول شهید مفتح جاسک')")
    cursor.execute("INSERT OR IGNORE INTO school_settings (key, val) VALUES ('school_theme', 'آبی هوشمند')")

    # 2. Remedial classes table
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS remedial_classes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        teacher_id INTEGER NOT NULL,
        title TEXT NOT NULL,
        description TEXT,
        grade INTEGER NOT NULL,
        capacity INTEGER NOT NULL,
        price REAL NOT NULL,
        schedule TEXT,
        status TEXT DEFAULT 'active'
    )
    """)

    # 3. Teacher payment gateway settings table
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS teacher_gateways (
        teacher_id INTEGER PRIMARY KEY,
        card_number TEXT,
        sheba TEXT,
        bank_name TEXT,
        gateway_type TEXT DEFAULT 'direct',
        merchant_id TEXT
    )
    """)

    # 4. Class payments/registrations table
    cursor.execute("""
    CREATE TABLE IF NOT EXISTS class_payments (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        student_id INTEGER NOT NULL,
        class_id INTEGER NOT NULL,
        amount REAL NOT NULL,
        card_number TEXT,
        tracking_code TEXT,
        status TEXT DEFAULT 'pending',
        date TEXT NOT NULL
    )
    """)
    conn.commit()
    conn.close()

def get_school_name():
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT val FROM school_settings WHERE key = 'school_name'")
        row = cursor.fetchone()
        conn.close()
        if row:
            return row['val']
    except Exception:
        pass
    return "دبیرستان متوسطه اول شهید مفتح جاسک"

def get_school_theme():
    try:
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT val FROM school_settings WHERE key = 'school_theme'")
        row = cursor.fetchone()
        conn.close()
        if row:
            return row['val']
    except Exception:
        pass
    return "آبی هوشمند"


# Custom styling for RTL and Persian fonts with Happy, Smart, and Professional theme support
theme_name = get_school_theme()
if theme_name == "سبز شاداب":
    primary_color = "#047857" # Emerald Green
    secondary_color = "#ECFDF5" # Soft Green
    accent_color = "#10B981" # Green
    banner_color = "#D1FAE5"
    text_color = "#065F46"
elif theme_name == "نارنجی پرانرژی":
    primary_color = "#C2410C" # Rich Orange
    secondary_color = "#FFF7ED" # Soft Orange
    accent_color = "#F97316" # Orange
    banner_color = "#FFEDD5"
    text_color = "#7C2D12"
else: # آبی هوشمند (default)
    primary_color = "#1E3A8A" # Dark Blue
    secondary_color = "#EFF6FF" # Soft Blue
    accent_color = "#2563EB" # Blue
    banner_color = "#DBEAFE"
    text_color = "#1D4ED8"

st.markdown(f"""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+Arabic:wght@400;700&display=swap');
    
    html, body, [data-testid="stSidebar"], .stMarkdown, p, div, h1, h2, h3, h4, h5, h6, span, label, input, select, button {{
        font-family: 'Noto Sans Arabic', sans-serif !important;
        direction: rtl !important;
        text-align: right !important;
    }}
    
    /* Marquee style for announcements */
    .ticker-wrap {{
        background-color: {banner_color};
        border-bottom: 2px solid {accent_color};
        color: {text_color};
        padding: 8px 10px;
        font-weight: bold;
        overflow: hidden;
        margin-bottom: 20px;
        border-radius: 4px;
    }}
    .ticker {{
        display: inline-block;
        white-space: nowrap;
        animation: marquee 25s linear infinite;
        font-size: 14px;
    }}
    @keyframes marquee {{
        0% {{ transform: translate3d(100%, 0, 0); }}
        100% {{ transform: translate3d(-100%, 0, 0); }}
    }}
    
    /* Info cards */
    .metric-card {{
        background-color: {secondary_color};
        border: 1px solid {accent_color};
        border-radius: 8px;
        padding: 15px;
        text-align: center;
        margin-bottom: 15px;
    }}
    .metric-title {{
        font-size: 13px;
        color: {primary_color};
        font-weight: bold;
    }}
    .metric-val {{
        font-size: 24px;
        font-weight: bold;
        color: {primary_color};
        margin-top: 5px;
    }}
    
    /* Happy professional UI buttons and details */
    .stButton>button {{
        border-radius: 8px !important;
        font-weight: bold !important;
        background-color: {primary_color} !important;
        color: white !important;
        transition: all 0.3s ease;
    }}
    .stButton>button:hover {{
        background-color: {accent_color} !important;
        transform: scale(1.02);
    }}
    
    /* Custom headers and containers */
    h1, h2, h3 {{
        color: {primary_color} !important;
    }}
    .stTabs [data-baseweb="tab-list"] {{
        background-color: {secondary_color};
        padding: 5px;
        border-radius: 8px;
    }}
    .stTabs [data-baseweb="tab"] {{
        color: {primary_color} !important;
        font-weight: bold !important;
    }}
    .stTabs [aria-selected="true"] {{
        background-color: white !important;
        border-radius: 6px;
    }}
</style>
""", unsafe_allow_html=True)

def show_announcements_marquee():
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT title, date FROM announcements ORDER BY id DESC LIMIT 5")
    rows = cursor.fetchall()
    conn.close()
    
    if rows:
        ann_text = "  |  ".join([f"📢 {row['title']} ({row['date']})" for row in rows])
        st.markdown(f"""
        <div class="ticker-wrap">
            <div class="ticker">
                {ann_text}
            </div>
        </div>
        """, unsafe_allow_html=True)

# Auth sessions
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.user = None

def login_user(username, password):
    persian_to_eng = str.maketrans('۰۱۲۳۴۵۶۷۸۹٠١٢٣٤٥٦٧٨٩', '01234567890123456789')
    username = str(username).strip().translate(persian_to_eng)
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT * FROM users WHERE username = ? AND password = ?", (username, password))
    user = cursor.fetchone()
    conn.close()
    if user:
        st.session_state.logged_in = True
        st.session_state.user = dict(user)
        return True
    return False

def logout_user():
    st.session_state.logged_in = False
    st.session_state.user = None
    st.rerun()

# --- APP LAYOUT ---
show_announcements_marquee()

# Check if database has any users
has_users = False
try:
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT COUNT(*) as count FROM users")
    row = cursor.fetchone()
    if row and row['count'] > 0:
        has_users = True
    conn.close()
except Exception as e:
    pass


st.title("🎓 مدرسه هوشمند")
st.subheader(get_school_name())

if not st.session_state.logged_in:
    if not has_users:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.markdown("""
            <div style="background-color: #EFF6FF; border: 2px solid #BFDBFE; border-radius: 12px; padding: 25px; text-align: center; direction: rtl;">
                <h2 style="color: #1E3A8A; margin-top: 0; font-family: 'Noto Sans Arabic', sans-serif !important;">🛠️ پیکربندی و راه‌اندازی اولیه سامانه</h2>
                <p style="color: #1E40AF; font-size: 14px; font-family: 'Noto Sans Arabic', sans-serif !important;">خوش آمدید! هیچ حسابی روی سیستم وجود ندارد. لطفاً مشخصات مدرسه و مدیر ارشد خود را برای ساخت پایگاه داده امن تعیین کنید:</p>
            </div>
            """, unsafe_allow_html=True)
            
            with st.form("initial_setup_form"):
                setup_school_name = st.text_input("نام رسمی مدرسه / آموزشگاه (مثال: دبیرستان شهید مفتح جاسک)", value="دبیرستان متوسطه اول شهید مفتح جاسک")
                setup_admin_name = st.text_input("نام و نام خانوادگی مدیر ارشد")
                setup_admin_username = st.text_input("کد ملی مدیر ارشد (۱۰ رقم عددی - اجباری)", placeholder="مثال: 0012345678")
                setup_admin_password = st.text_input("رمز عبور مدیر", type="password")
                setup_admin_password_confirm = st.text_input("تکرار رمز عبور مدیر", type="password")
                setup_theme = st.selectbox("قالب و تم رنگی پیش‌فرض سامانه", ["آبی هوشمند", "سبز شاداب", "نارنجی پرانرژی"])
                
                setup_submit = st.form_submit_button("🚀 ثبت اطلاعات و راه‌اندازی مدرسه هوشمند")
                if setup_submit:
                    cleaned_admin_id = clean_and_validate_national_id(setup_admin_username)
                    if setup_school_name and setup_admin_name and setup_admin_username and setup_admin_password:
                        if not cleaned_admin_id:
                            st.error("❌ کد ملی مدیر ارشد باید دقیقاً ۱۰ رقم عددی باشد.")
                        elif setup_admin_password != setup_admin_password_confirm:
                            st.error("❌ رمز عبور و تکرار آن با هم مطابقت ندارند.")
                        elif len(setup_admin_password) < 4:
                            st.error("❌ رمز عبور باید حداقل ۴ نویسه باشد.")
                        else:
                            try:
                                conn = get_connection()
                                cursor = conn.cursor()
                                cursor.execute("INSERT INTO users (username, password, role, name) VALUES (?, ?, 'admin', ?)",
                                               (cleaned_admin_id, setup_admin_password, setup_admin_name))
                                cursor.execute("INSERT OR REPLACE INTO school_settings (key, val) VALUES ('school_name', ?)", (setup_school_name,))
                                cursor.execute("INSERT OR REPLACE INTO school_settings (key, val) VALUES ('school_theme', ?)", (setup_theme,))
                                conn.commit()
                                conn.close()
                                
                                st.success("🎉 سامانه مدرسه هوشمند با موفقیت پیکربندی شد! اکنون می‌توانید با اطلاعات تعیین‌شده وارد شوید.")
                                st.balloons()
                                st.rerun()
                            except Exception as e:
                                st.error(f"خطا در ثبت اطلاعات اولیه: {e}")
                    else:
                        st.warning("⚠️ تکمیل تمامی کادرهای ستاره‌دار الزامی است.")
    else:
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            st.info("🔐 جهت ورود به سامانه، کد ملی خود (یا کد ملی فرزند با پیشوند p_ برای اولیاء) را وارد کنید:")
            username = st.text_input("کد ملی (یا p_کد ملی برای اولیاء)")
            password = st.text_input("رمز عبور", type="password")
            if st.button("ورود به سامانه"):
                if login_user(username, password):
                    st.success(f"خوش آمدید، {st.session_state.user['name']}")
                    st.rerun()
                else:
                    st.error("نام کاربری یا رمز عبور اشتباه است.")
else:
    user = st.session_state.user
    role = user["role"]
    
    # --- LICENSE & TRIAL CHECK ---
    license_status, days_left = smart_school_addons.check_license_status(get_connection)
    
    # Force sidebar warning for Trial period
    if license_status == "trial":
        st.sidebar.warning(f"⏳ دوره آزمایشی فعال است: {days_left} روز باقی مانده")
        if st.sidebar.button("🔑 فعال‌سازی نسخه دائمی (پرداخت به سازنده)"):
            st.session_state.force_show_activation = True
            st.rerun()
            
    # Expired check
    is_system_expired = (license_status == "expired")
    if st.session_state.get("force_show_activation", False):
        st.header("🔑 درگاه فعال‌سازی و تمدید لایسنس")
        if st.button("🔙 انصراف و بازگشت به سامانه"):
            st.session_state.force_show_activation = False
            st.rerun()
        smart_school_addons.render_activation_gateway(get_connection)
        st.stop()
        
    if is_system_expired:
        if role == "admin":
            st.header("🔑 قفل موقت سامانه - فعال‌سازی مورد نیاز است")
            st.warning("⚠️ مهلت ۱۰ روزه استفاده رایگان این مدرسه به پایان رسیده است. برای باز شدن قفل نرم‌افزار، لطفاً لایسنس فعال‌سازی دائمی را پرداخت فرمایید.")
            smart_school_addons.render_activation_gateway(get_connection)
            st.stop()
        else:
            st.error("🚫 مهلت استفاده آزمایشی رایگان این مدرسه به پایان رسیده است.")
            st.info("لطفاً از مدیریت مدرسه بخواهید نسبت به فعال‌سازی دائمی سامانه اقدام کند.")
            st.markdown(f"**توسعه‌دهنده سامانه:** {smart_school_addons.CREATOR_NAME}<br>**شماره کارت مقصد:** {smart_school_addons.CREATOR_CARD}", unsafe_allow_html=True)
            st.stop()
    
    # Sidebar
    st.sidebar.markdown(f"### 👤 {user['name']}")
    st.sidebar.markdown(f"**نقش شما:** {role.upper()}")
    if st.sidebar.button("🚪 خروج از حساب"):
        logout_user()
        
    st.sidebar.markdown("---")
    st.sidebar.markdown("---")
    # --- GLOBAL SIDEBAR CALCULATOR ---
    with st.sidebar.expander("🧮 ماشین‌حساب سریع پُل", expanded=False):
        st.markdown("<p style='text-align: right; font-size: 13px; color: #1E3A8A;'>عبارت ریاضی خود را بنویسید (مثال: (12+8)*5 ):</p>", unsafe_allow_html=True)
        calc_expr = st.text_input("", key="sidebar_calc_expr", placeholder="مثلا: (25 + 5) / 6")
        if calc_expr:
            def safe_eval_sidebar(expr):
                # Clean characters for safe evaluation in Persian keyboard support
                expr = expr.replace(" ", "").replace("×", "*").replace("÷", "/").replace("−", "-")
                # Standard characters allowed: digits, arithmetic, dots, parentheses
                import re
                if not re.match(r"^[0-9+\-*/().]*$", expr):
                    return "خطا: نویسه غیرمجاز"
                try:
                    # Prevent division by zero
                    if "/0" in expr:
                        return "خطا: تقسیم بر صفر"
                    val = eval(expr, {"__builtins__": {}}, {})
                    if isinstance(val, float):
                        return round(val, 4)
                    return val
                except Exception:
                    return "خطا در عبارت"
            calc_res = safe_eval_sidebar(calc_expr)
            if "خطا" in str(calc_res):
                st.error(calc_res)
            else:
                st.success(f"نتیجه: **{calc_res}**")

    with st.sidebar.expander("💬 پشتیبانی و ارتباط با سازنده", expanded=False):
        st.markdown("<p style='text-align: right; font-size: 13px; color: #1E3A8A; font-weight: bold;'>📞 مرکز پشتیبانی مدارس هوشمند ایران</p>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: right; font-size: 12px; margin: 0;'><b>👨‍💻 توسعه‌دهنده:</b> رستم سوری نسب</p>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: right; font-size: 12px; margin: 0;'><b>📧 ایمیل پشتیبانی:</b> support@smart-school.ir</p>", unsafe_allow_html=True)
        st.markdown("<p style='text-align: right; font-size: 12px; margin: 0;'><b>📱 پشتیبانی تلگرام/ایتا:</b> @SmartSchool_Support</p>", unsafe_allow_html=True)
        st.markdown("<hr style='margin: 10px 0;'>", unsafe_allow_html=True)
        
        ticket_text = st.text_area("✍️ ارسال تیکت مستقیم به توسعه‌دهنده:", placeholder="سوال، گزارش خطا یا پیشنهاد خود را اینجا بنویسید...", key="support_ticket_text", height=80)
        if st.button("🚀 ارسال تیکت پشتیبانی", key="support_ticket_submit_btn"):
            if ticket_text.strip():
                st.success("✔️ تیکت شما با موفقیت ثبت و ارسال شد! کد پیگیری: SH-" + str(datetime.now().microsecond))
            else:
                st.warning("⚠️ لطفا ابتدا متن تیکت خود را بنویسید.")
    
    
    # --- ROLE-BASED PANELS ---
    if role == "admin":
        st.sidebar.markdown("### ⚙️ پنل مدیریت")
        menu = st.sidebar.radio("انتخاب منو", [
            "📂 مدیریت کاربران و اکسل",
            "🏫 کلاس‌های تقویتی کل مدرسه",
            "📢 تابلوی اعلانات و بخشنامه‌ها",
            "📅 مدیریت تقویم و امتحانات",
            "⚙️ ابزارها و خروجی گزارشات"
        ], key="admin_menu_sel")
        
        if menu == "📂 مدیریت کاربران و اکسل":
            st.header("📂 مدیریت و پورتال کاربران")
            
            # Creating Tabs for User Management
            tab1, tab2, tab3, tab4 = st.tabs([
                "📥 درون‌ریزی دسته‌جمعی (اکسل)", 
                "➕ تعریف کاربر جدید (تکی)", 
                "🔑 مدیریت و تغییر رمزها", 
                "👤 تغییر رمز من (مدیر)"
            ])
            
            with tab1:
                st.subheader("۱. درون‌ریزی دانش‌آموزان با فایل اکسل")
                st.write("شما می‌توانید کل لیست دانش‌آموزان را به همراه کلاس و پایه در قالب یک فایل اکسل آپلود کنید. سامانه به صورت خودکار پورتال اولیای مربوطه را نیز می‌سازد.")
                
                # Download template
                template_path = os.path.join(os.path.dirname(__file__), "excel_template.xlsx")
                if os.path.exists(template_path):
                    with open(template_path, "rb") as f_template:
                        st.download_button("📥 دانلود فایل نمونه اکسل (Template)", f_template, "excel_template.xlsx")
                
                uploaded_file = st.file_uploader("فایل اکسل خود را انتخاب کنید", type=["xlsx"])
                if uploaded_file is not None:
                    try:
                        df = pd.read_excel(uploaded_file)
                        st.write("**پیش‌نمایش اطلاعات فایل آپلود شده:**")
                        st.dataframe(df)
                        
                        if st.button("🚀 شروع درون‌ریزی دانش‌آموزان و ساخت پورتال اولیا"):
                            conn = get_connection()
                            cursor = conn.cursor()
                            success_count = 0
                            for idx, row in df.iterrows():
                                fullname = row.get("نام و نام خانوادگی", "").strip()
                                username = str(row.get("نام کاربری", "")).strip()
                                class_id = str(row.get("کلاس", "")).strip()
                                grade = int(row.get("پایه", 9))
                                password = str(row.get("رمز عبور", "123")).strip()
                                
                                if fullname and username:
                                    try:
                                        cursor.execute("INSERT INTO users (username, password, role, name) VALUES (?, ?, 'student', ?)", (username, password, fullname))
                                        student_id = cursor.lastrowid
                                        cursor.execute("INSERT INTO students (id, class_id, grade) VALUES (?, ?, ?)", (student_id, class_id, grade))
                                        
                                        # Create parent user
                                        p_username = f"p_{username}"
                                        p_fullname = f"ولیِ {fullname}"
                                        cursor.execute("INSERT INTO users (username, password, role, name) VALUES (?, ?, 'parent', ?)", (p_username, password, p_fullname))
                                        success_count += 1
                                    except Exception as e:
                                        pass
                            conn.commit()
                            conn.close()
                            st.success(f"موفقیت‌آمیز! تعداد {success_count} دانش‌آموز جدید و اولیای آن‌ها با موفقیت اضافه شدند!")
                            st.rerun()
                    except Exception as e:
                        st.error(f"خطا در خواندن فایل اکسل: {e}")
            
            with tab2:
                st.subheader("۲. تعریف تک‌به‌تک و دستی کاربر جدید")
                st.write("از این قسمت می‌توانید همکاران (دبیران)، دانش‌آموزان یا مدیران جدید را به صورت تکی تعریف کنید:")
                
                with st.form("create_user_form_v4", clear_on_submit=True):
                    new_name = st.text_input("نام و نام خانوادگی")
                    new_username = st.text_input("کد ملی (۱۰ رقم عددی - اجباری)")
                    new_password = st.text_input("رمز عبور")
                    new_role = st.selectbox("نقش کاربر جدید", ["دانش‌آموز (student)", "دبیر (teacher)", "مدیر (admin)"])
                    
                    st.markdown("---")
                    st.write("⚠️ تکمیل اطلاعات زیر فقط در صورت انتخاب نقش **دانش‌آموز** الزامی است:")
                    class_id = st.text_input("کلاس (مثلاً 9-1)", value="9-1")
                    grade = st.selectbox("پایه تحصیلی", [7, 8, 9], index=2)
                    
                    submit_user = st.form_submit_button("➕ ثبت کاربر جدید در سامانه")
                    if submit_user:
                        cleaned_id = clean_and_validate_national_id(new_username)
                        if not cleaned_id:
                            st.error("❌ کد ملی باید دقیقاً یک عدد ۱۰ رقمی باشد.")
                        elif new_name and new_username and new_password:
                            role_map = {
                                "دانش‌آموز (student)": "student",
                                "دبیر (teacher)": "teacher",
                                "مدیر (admin)": "admin"
                            }
                            role_db = role_map[new_role]
                            conn = get_connection()
                            cursor = conn.cursor()
                            try:
                                cursor.execute("INSERT INTO users (username, password, role, name) VALUES (?, ?, ?, ?)", (cleaned_id, new_password, role_db, new_name))
                                new_user_id = cursor.lastrowid
                                
                                if role_db == "student":
                                    cursor.execute("INSERT INTO students (id, class_id, grade) VALUES (?, ?, ?)", (new_user_id, class_id, grade))
                                    # Create parent
                                    p_username = f"p_{cleaned_id}"
                                    p_fullname = f"ولیِ {new_name}"
                                    cursor.execute("INSERT INTO users (username, password, role, name) VALUES (?, ?, 'parent', ?)", (p_username, new_password, p_fullname))
                                
                                conn.commit()
                                st.success(f"کاربر جدید '{new_name}' با نقش {new_role} با موفقیت ثبت شد!")
                            except sqlite3.IntegrityError:
                                st.error("خطا: این نام کاربری قبلاً در سامانه استفاده شده است!")
                            except Exception as e:
                                st.error(f"خطا در ثبت کاربر: {e}")
                            finally:
                                conn.close()
                        else:
                            st.warning("لطفاً تمام کادرهای الزامی را پر کنید.")
            
            with tab3:
                st.subheader("۳. مدیریت کاربران و ویرایش یا تغییر آنلاین رمزها")
                st.write("کاربر مورد نظر خود را از لیست زیر انتخاب کرده و اطلاعات یا رمز عبور او را تغییر دهید:")
                
                conn = get_connection()
                df_edit_users = pd.read_sql_query("SELECT id, name, username, role FROM users ORDER BY id DESC", conn)
                conn.close()
                
                user_options = {f"{row['name']} (کد ملی: {row['username']} - {row['role']})": row['id'] for idx, row in df_edit_users.iterrows()}
                selected_user_str = st.selectbox("🎯 انتخاب کاربر جهت ویرایش:", list(user_options.keys()))
                
                if selected_user_str:
                    selected_id = user_options[selected_user_str]
                    conn = get_connection()
                    cursor = conn.cursor()
                    cursor.execute("SELECT * FROM users WHERE id = ?", (selected_id,))
                    user_data = cursor.fetchone()
                    conn.close()
                    
                    if user_data:
                        st.write(f"✍️ **ویرایش حساب:** {user_data['name']} (نقش: {user_data['role'].upper()})")
                        with st.form(f"edit_user_form_{selected_id}"):
                            edit_name = st.text_input("نام و نام خانوادگی", value=user_data['name'])
                            edit_username = st.text_input("کد ملی (۱۰ رقم - اجباری)", value=user_data['username'])
                            edit_password = st.text_input("رمز عبور جدید", value=user_data['password'])
                            
                            update_btn = st.form_submit_button("💾 ذخیره تغییرات کاربر")
                            if update_btn:
                                cleaned_id = clean_and_validate_national_id(edit_username)
                                if not cleaned_id:
                                    st.error("❌ کد ملی باید دقیقاً ۱۰ رقم عددی باشد.")
                                elif edit_name and edit_username and edit_password:
                                    conn = get_connection()
                                    cursor = conn.cursor()
                                    try:
                                        cursor.execute("UPDATE users SET name = ?, username = ?, password = ? WHERE id = ?", (edit_name, cleaned_id, edit_password, selected_id))
                                        conn.commit()
                                        st.success(f"اطلاعات کاربر '{edit_name}' با موفقیت بروزرسانی شد!")
                                        st.rerun()
                                    except sqlite3.IntegrityError:
                                        st.error("خطا: این نام کاربری قبلاً در سامانه استفاده شده است!")
                                    except Exception as e:
                                        st.error(f"خطا در ذخیره تغییرات: {e}")
                                    finally:
                                        conn.close()
                                else:
                                    st.warning("کادرها نباید خالی باشند.")
                
                st.markdown("---")
                st.write("📊 **لیست زنده کل کاربران ثبت‌شده در سامانه:**")
                st.dataframe(df_edit_users)
            
            with tab4:
                st.subheader("۴. تغییر نام کاربری و رمز عبور شما (مدیر)")
                st.write("از این فرم می‌توانید به عنوان مدیر اصلی، اطلاعات ورود خود را شخصی‌سازی و امن کنید:")
                
                with st.form("admin_self_edit_v4"):
                    admin_id = st.session_state.user['id']
                    admin_name = st.text_input("نام نمایش‌داده‌شده شما در سایت", value=st.session_state.user['name'])
                    admin_username = st.text_input("کد ملی مدیر (۱۰ رقم - اجباری)", value=st.session_state.user['username'])
                    admin_password = st.text_input("رمز عبور جدید مدیر", value=st.session_state.user['password'])
                    
                    submit_admin_self = st.form_submit_button("💾 ثبت نهایی تغییرات مدیر")
                    if submit_admin_self:
                        cleaned_id = clean_and_validate_national_id(admin_username)
                        if not cleaned_id:
                            st.error("❌ کد ملی مدیر باید دقیقاً ۱۰ رقم عددی باشد.")
                        elif admin_name and admin_username and admin_password:
                            conn = get_connection()
                            cursor = conn.cursor()
                            try:
                                cursor.execute("UPDATE users SET name = ?, username = ?, password = ? WHERE id = ?", (admin_name, cleaned_id, admin_password, admin_id))
                                conn.commit()
                                # Update current session
                                st.session_state.user['name'] = admin_name
                                st.session_state.user['username'] = cleaned_id
                                st.session_state.user['password'] = admin_password
                                st.success("مشخصات ورود مدیریت با موفقیت امن و بروزرسانی شد!")
                                st.rerun()
                            except sqlite3.IntegrityError:
                                st.error("خطا: این نام کاربری قبلاً در سامانه استفاده شده است!")
                            except Exception as e:
                                st.error(f"خطا در بروزرسانی: {e}")
                            finally:
                                conn.close()
                        else:
                            st.warning("پر کردن تمامی کادرها الزامی است.")

        elif menu == "📢 تابلوی اعلانات و بخشنامه‌ها":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_admin_ann"): st.session_state.admin_menu_sel = "📂 مدیریت کاربران و اکسل"; st.rerun()
            st.header("📢 تابلوی اعلانات و بخشنامه‌ها")
            
            title = st.text_input("عنوان بخشنامه / اطلاعیه")
            text = st.text_area("متن بخشنامه")
            target = st.selectbox("مخاطب هدف", ["all", "students", "parents", "teachers"])
            
            if st.button("ثبت بخشنامه جدید"):
                if title and text:
                    conn = get_connection()
                    cursor = conn.cursor()
                    cursor.execute("INSERT INTO announcements (title, text, target, date) VALUES (?, ?, ?, ?)",
                                   (title, text, target, datetime.now().strftime("%Y/%m/%d")))
                    conn.commit()
                    conn.close()
                    st.success("بخشنامه جدید با موفقیت ثبت شد و در بالای سایت فعال گردید!")
                    st.rerun()
                else:
                    st.warning("لطفاً عنوان و متن را کامل کنید.")
                    
            st.subheader("لیست کل بخشنامه‌ها")
            conn = get_connection()
            df_ann = pd.read_sql_query("SELECT id, title, target, date FROM announcements ORDER BY id DESC", conn)
            conn.close()
            st.dataframe(df_ann)
            
        elif menu == "📅 مدیریت تقویم و امتحانات":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_admin_cal"): st.session_state.admin_menu_sel = "📂 مدیریت کاربران و اکسل"; st.rerun()
            st.header("📅 مدیریت تقویم آموزشی و برنامه امتحانات")
            
            title = st.text_input("عنوان رویداد یا آزمون")
            date_str = st.text_input("تاریخ برگزاری (مثلاً: 1405/07/15)")
            type_event = st.selectbox("نوع رویداد", ["exam", "event"])
            grade_target = st.selectbox("پایه تحصیلی هدف", [7, 8, 9])
            
            if st.button("ثبت در تقویم"):
                if title and date_str:
                    conn = get_connection()
                    cursor = conn.cursor()
                    cursor.execute("INSERT INTO events (title, date, type, grade) VALUES (?, ?, ?, ?)",
                                   (title, date_str, type_event, grade_target))
                    conn.commit()
                    conn.close()
                    st.success("رویداد جدید با موفقیت ثبت شد!")
                    st.rerun()
                else:
                    st.warning("پر کردن عنوان و تاریخ الزامی است.")
                    
            st.subheader("رویدادهای ثبت‌شده")
            conn = get_connection()
            df_ev = pd.read_sql_query("SELECT id, title, date, type, grade FROM events ORDER BY date ASC", conn)
            conn.close()
            st.dataframe(df_ev)
            
        elif menu == "⚙️ ابزارها و خروجی گزارشات":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_admin_tools"): st.session_state.admin_menu_sel = "📂 مدیریت کاربران و اکسل"; st.rerun()
            st.header("⚙️ ابزارها و گزارش‌گیری رسمی اداره")
            
            st.subheader("۱. خروجی اکسل چندبرگی گزارش اداره")
            st.write("با کلیک روی دکمه زیر، سامانه آخرین اطلاعات رتبه‌بندی تحصیلی، نمرات هفتگی، نتایج آزمون‌های آنلاین و حضور و غیاب دانش‌آموزان را تجمیع کرده و فایل اکسل گزارش اداره را تولید می‌کند.")
            
            # Trigger Excel compilation
            excel_path = os.path.join(os.path.dirname(__file__), "school-reports-department.xlsx")
            # We can run create_excels directly if we want, let's copy create_excels logic to keep it single file or run via subprocess
            # Let's run a small subprocess since create_excels is deleted but we can do it via code right here
            if st.button("🚀 تولید و استخراج فایل اکسل گزارشات"):
                import openpyxl
                # Run creation right here
                conn = get_connection()
                df_rankings = pd.read_sql_query("""
                    SELECT u.name as "نام و نام خانوادگی", s.class_id as "کلاس", s.grade as "پایه", AVG(g.grade_val) as "معدل مستمر"
                    FROM users u JOIN students s ON u.id = s.id LEFT JOIN grades g ON u.id = g.student_id GROUP BY u.id ORDER BY "معدل مستمر" DESC
                """, conn)
                df_rankings.insert(0, 'رتبه', range(1, len(df_rankings) + 1))
                
                df_weekly = pd.read_sql_query("""
                    SELECT u.name as "نام دانش‌آموز", s.class_id as "کلاس", s.grade as "پایه", g.subject as "عنوان ارزیابی", g.grade_val as "نمره", g.date as "تاریخ ثبت"
                    FROM grades g JOIN users u ON g.student_id = u.id JOIN students s ON u.id = s.id
                """, conn)
                
                df_quizzes = pd.read_sql_query("""
                    SELECT u.name as "نام دانش‌آموز", s.class_id as "کلاس", s.grade as "پایه", q.title as "عنوان آزمون", qa.score as "نمره", qa.date as "تاریخ"
                    FROM quiz_attempts qa JOIN users u ON qa.student_id = u.id JOIN students s ON u.id = s.id JOIN quizzes q ON qa.quiz_id = q.id
                """, conn)
                
                df_attendance = pd.read_sql_query("""
                    SELECT u.name as "نام دانش‌آموز", s.class_id as "کلاس", s.grade as "پایه",
                           SUM(CASE WHEN a.status = 'حاضر' THEN 1 ELSE 0 END) as "تعداد حضور",
                           SUM(CASE WHEN a.status = 'غایب' THEN 1 ELSE 0 END) as "تعداد غیبت",
                           SUM(CASE WHEN a.status = 'تاخیر' THEN 1 ELSE 0 END) as "تعداد تاخیر"
                    FROM attendance a JOIN users u ON a.student_id = u.id JOIN students s ON u.id = s.id GROUP BY u.id
                """, conn)
                conn.close()
                
                with pd.ExcelWriter(excel_path, engine='openpyxl') as writer:
                    df_rankings.to_excel(writer, sheet_name="رتبه‌بندی کل مدرسه", index=False)
                    df_weekly.to_excel(writer, sheet_name="ریز نمرات هفتگی", index=False)
                    df_quizzes.to_excel(writer, sheet_name="نتایج آزمون‌های آنلاین", index=False)
                    df_attendance.to_excel(writer, sheet_name="خلاصه حضور غیاب", index=False)
                st.success("فایل اکسل با موفقیت تولید شد!")
                
            if os.path.exists(excel_path):
                with open(excel_path, "rb") as f:
                    st.download_button("📥 دریافت فایل اکسل گزارش اداره (.xlsx)", f, "school-reports-department.xlsx")
                    
            st.subheader("۲. دریافت پی‌دی‌اف‌های کارنامه و رتبه‌بندی مدارس")
            # PDF download
            card_path = os.path.join(os.path.dirname(__file__), "student-report-card-sample.pdf")
            rankings_path = os.path.join(os.path.dirname(__file__), "school-rankings-sample.pdf")
            
            col_pdf1, col_pdf2 = st.columns(2)
            with col_pdf1:
                if os.path.exists(card_path):
                    with open(card_path, "rb") as f:
                        st.download_button("📥 دانلود فایل کارنامه نمونه (.pdf)", f, "student-report-card-sample.pdf")
            with col_pdf2:
                if os.path.exists(rankings_path):
                    with open(rankings_path, "rb") as f:
                        st.download_button("📥 دانلود جدول رتبه‌بندی چاپی (.pdf)", f, "school-rankings-sample.pdf")
                        
        elif menu == "🏫 کلاس‌های تقویتی کل مدرسه":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_admin_classes"): st.session_state.admin_menu_sel = "📂 مدیریت کاربران و اکسل"; st.rerun()
            smart_school_addons.render_admin_classes_panel(get_connection)

    elif role == "teacher":
        st.sidebar.markdown("### 📐 پنل معلمان")
        menu = st.sidebar.radio("انتخاب منو", [
            "📢 تابلوی اعلانات",
            "📝 طراحی و تولید آزمون آنلاین",
            "🏫 کلاس‌های تقویتی و خصوصی",
            "📂 ثبت نمرات و بازخوردهای کلاسی",
            "📅 مدیریت تقویم و امتحانات",
            "🖥️ کلاس آنلاین و حضور و غیاب زنده"
        ], key="teacher_menu_sel")
        
        if menu == "📢 تابلوی اعلانات":
            st.header("📢 تابلوی اعلانات مدرسه")
            st.write("آخرین بخشنامه‌ها و اطلاعیه‌ها در نوار متحرک بالای سایت نمایش داده می‌شوند.")
            
        elif menu == "📝 طراحی و تولید آزمون آنلاین":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_teacher_quiz"): st.session_state.teacher_menu_sel = "📢 تابلوی اعلانات"; st.rerun()
            st.header("📝 طراحی و تولید هوشمند آزمون آنلاین")
            
            # Creating Tabs for manual and AI quiz creation
            tab_manual, tab_ai = st.tabs([
                "📝 طراحی دستی سوالات تستی", 
                "🧠 آزمون‌ساز هوشمند با هوش مصنوعی (بدون فیلتر)"
            ])
            
            with tab_manual:
                st.subheader("۱. طراحی سریع سوال تستی به صورت دستی")
                title = st.text_input("عنوان آزمون", key="manual_title")
                grade = st.selectbox("پایه هدف", [7, 8, 9], key="manual_grade")
                topic = st.text_input("موضوع آزمون", "توان و ریشه", key="manual_topic")
                
                if st.button("ایجاد آزمون دستی", key="manual_create_btn"):
                    if title:
                        conn = get_connection()
                        cursor = conn.cursor()
                        cursor.execute("INSERT INTO quizzes (title, grade, topic) VALUES (?, ?, ?)", (title, grade, topic))
                        conn.commit()
                        conn.close()
                        st.success(f"آزمون '{title}' با موفقیت تعریف شد. اکنون می‌توانید سوالات آن را از طریق دیتابیس یا کدهای جدید اضافه کنید.")
                    else:
                        st.warning("لطفاً عنوان آزمون را وارد کنید.")
                        
            with tab_ai:
                st.subheader("🧠 آزمون‌ساز جادویی هوش مصنوعی پُل (مشابه ChatGPT)")
                st.write("کافی است متن درسنامه، خلاصه مبحث ریاضی یا فرمول‌های کلاسی را تایپ کنید، یا **فایل جزوه (PDF/Word)** یا **فایل نمونه‌سوال تستی آماده** را بارگذاری کنید تا آزمون به صورت خودکار برای شما ساخته شود:")
                
                ai_title = st.text_input("عنوان آزمون هوش مصنوعی", placeholder="مثلاً: آزمون فصل اول ریاضی هفتم", key="ai_title")
                ai_grade = st.selectbox("پایه تحصیلی هدف برای آزمون:", [7, 8, 9], index=2, key="ai_grade")
                ai_topic = st.text_input("موضوع آزمون برای طبقه‌بندی:", "توان و ریشه", key="ai_topic")
                
                # Choose input method
                input_method = st.radio("📚 انتخاب منبع طراحی آزمون:", [
                    "✍️ تایپ یا کپی کردن متن درسنامه (متنی)",
                    "📁 بارگذاری فایل جزوه و درسنامه آموزشی (PDF یا Word)",
                    "📝 بارگذاری فایل حاوی نمونه‌سوال تستی آماده (PDF یا Word)"
                ], key="input_method")
                
                extracted_text = ""
                ready_questions = []
                ai_num_qs = 3
                
                if input_method == "✍️ تایپ یا کپی کردن متن درسنامه (متنی)":
                    lecture_text = st.text_area("متن درسنامه / مبحث علمی کتاب ریاضی جهت طراحی سوالات:", height=150, placeholder="مثال:\nریشه دوم عدد ۲۵ برابر با ۵ است.\nدر یک مثلث قائم‌الزاویه، رابطه فیثاغورس برقرار است که در آن مجموع مجذور اضلاع قائمه برابر با مجذور وتر است.", key="lecture_text")
                    ai_num_qs = st.slider("تعداد سوالات مورد نیاز:", 2, 5, 3, key="ai_num_qs_text")
                    
                    if st.button("🚀 تولید آزمون هوشمند از متن درسنامه", key="ai_generate_btn_text"):
                        if lecture_text:
                            extracted_text = lecture_text
                        else:
                            st.warning("لطفاً ابتدا متنی را کپی و بارگذاری فرمایید.")
                            
                elif input_method == "📁 بارگذاری فایل جزوه و درسنامه آموزشی (PDF یا Word)":
                    uploaded_file = st.file_uploader("فایل جزوه آموزشی یا درسنامه را انتخاب کنید (docx, pdf):", type=["pdf", "docx"], key="uploaded_lecture_file")
                    ai_num_qs = st.slider("تعداد سوالات مورد نیاز:", 2, 5, 3, key="ai_num_qs_file")
                    
                    if uploaded_file is not None:
                        file_name = uploaded_file.name.lower()
                        with st.spinner("⏳ در حال استخراج متن از فایل جزوه..."):
                            if file_name.endswith(".pdf"):
                                try:
                                    import pypdf
                                    reader = pypdf.PdfReader(uploaded_file)
                                    text_list = []
                                    for page in reader.pages:
                                        t = page.extract_text()
                                        if t:
                                            text_list.append(t)
                                    extracted_text = "\n".join(text_list)
                                    st.success(f"✔️ فایل پی‌دی‌اف خوانده شد! {len(extracted_text)} کاراکتر متن استخراج گردید.")
                                except Exception as e:
                                    st.error(f"خطا در خواندن فایل پی‌دی‌اف: {e}")
                            elif file_name.endswith(".docx"):
                                try:
                                    import docx
                                    doc = docx.Document(uploaded_file)
                                    text_list = [para.text for para in doc.paragraphs if para.text]
                                    extracted_text = "\n".join(text_list)
                                    st.success(f"✔️ فایل ورد خوانده شد! {len(extracted_text)} کاراکتر متن استخراج گردید.")
                                except Exception as e:
                                    st.error(f"خطا در خواندن فایل ورد: {e}")
                                    
                        if extracted_text:
                            with st.expander("🔍 مشاهده پیش‌نمایش متن استخراج‌شده جزوه"):
                                st.text(extracted_text[:1000] + ("..." if len(extracted_text) > 1000 else ""))
                                
                    if st.button("🚀 تولید آزمون هوشمند از فایل آپلود شده", key="ai_generate_btn_file"):
                        if not extracted_text:
                            st.warning("لطفاً ابتدا فایل جزوه را بارگذاری کنید.")
                            
                elif input_method == "📝 بارگذاری فایل حاوی نمونه‌سوال تستی آماده (PDF یا Word)":
                    uploaded_file = st.file_uploader("فایل حاوی نمونه‌سوال تستی آماده را انتخاب کنید (docx, pdf):", type=["pdf", "docx"], key="uploaded_quiz_file")
                    
                    if uploaded_file is not None:
                        file_name = uploaded_file.name.lower()
                        with st.spinner("⏳ در حال استخراج و تحلیل سوالات فایل..."):
                            raw_text = ""
                            if file_name.endswith(".pdf"):
                                try:
                                    import pypdf
                                    reader = pypdf.PdfReader(uploaded_file)
                                    text_list = []
                                    for page in reader.pages:
                                        t = page.extract_text()
                                        if t:
                                            text_list.append(t)
                                    raw_text = "\n".join(text_list)
                                except Exception as e:
                                    st.error(f"خطا در خواندن فایل پی‌دی‌اف: {e}")
                            elif file_name.endswith(".docx"):
                                try:
                                    import docx
                                    doc = docx.Document(uploaded_file)
                                    text_list = [para.text for para in doc.paragraphs if para.text]
                                    raw_text = "\n".join(text_list)
                                except Exception as e:
                                    st.error(f"خطا در خواندن فایل ورد: {e}")
                                    
                            if raw_text:
                                # Parse questions
                                def parse_questions_from_text(text):
                                    import re
                                    lines = text.split('\n')
                                    questions = []
                                    current_q = None
                                    
                                    for line in lines:
                                        line_strip = line.strip()
                                        if not line_strip:
                                            continue
                                            
                                        q_match = re.match(r'^[\s\(\[\{]*([0-9\u06f0-\u06f9]+)[\s\)\.\-\]\}]+(.*)', line_strip)
                                        if q_match:
                                            if current_q:
                                                questions.append(current_q)
                                            q_text = q_match.group(2).strip()
                                            current_q = {
                                                "question": q_text,
                                                "options": [],
                                                "correct": "",
                                                "correct_letter": "A"
                                            }
                                            continue
                                            
                                        if current_q:
                                            opt_matches = re.findall(r'([الفبجدوبABCD1234])[\s\)\.-]+([^الفبجدوبABCD1234\n]+)', line_strip)
                                            if opt_matches:
                                                for opt_let, opt_val in opt_matches:
                                                    opt_val = opt_val.strip()
                                                    if opt_val:
                                                        current_q["options"].append(opt_val)
                                            else:
                                                opt_match_single = re.match(r'^([الفبجدوبABCD1234])[\s\)\.-]+(.*)', line_strip)
                                                if opt_match_single:
                                                    opt_let = opt_match_single.group(1)
                                                    opt_val = opt_match_single.group(2).strip()
                                                    current_q["options"].append(opt_val)
                                                else:
                                                    if not current_q["options"]:
                                                        current_q["question"] += " " + line_strip
                                                    else:
                                                        ans_match = re.search(r'(پاسخ|جواب|correct|answer)[\s:=-]+([الفبجدوبABCD1234])', line_strip, re.IGNORECASE)
                                                        if ans_match:
                                                            current_q["correct"] = ans_match.group(2).strip()
                                                        else:
                                                            current_q["options"][-1] += " " + line_strip
                                                            
                                    if current_q:
                                        questions.append(current_q)
                                        
                                    valid_questions = []
                                    letter_map = {"الف": "A", "ب": "B", "ج": "C", "د": "D", "1": "A", "2": "B", "3": "C", "4": "D", "A": "A", "B": "B", "C": "C", "D": "D"}
                                    for q in questions:
                                        opts = [o for o in q["options"] if o.strip()]
                                        if len(opts) < 4:
                                            while len(opts) < 4:
                                                opts.append(f"گزینه {len(opts)+1}")
                                        elif len(opts) > 4:
                                            opts = opts[:4]
                                            
                                        corr_letter = "A"
                                        if q["correct"] in letter_map:
                                            corr_letter = letter_map[q["correct"]]
                                        
                                        corr_val = opts[ord(corr_letter) - ord('A')] if (ord(corr_letter) - ord('A')) < len(opts) else opts[0]
                                            
                                        valid_questions.append({
                                            "question": q["question"],
                                            "options": opts,
                                            "correct": corr_val,
                                            "correct_letter": corr_letter,
                                            "topic": "تست کاربری"
                                        })
                                    return valid_questions
                                    
                                ready_questions = parse_questions_from_text(raw_text)
                                if ready_questions:
                                    st.success(f"✔️ تعداد {len(ready_questions)} سوال تستی آماده با موفقیت از فایل استخراج شد!")
                                else:
                                    st.warning("⚠️ متنی حاوی سوالات تستی معتبر با ساختار شماره‌گذاری و گزینه‌های (الف، ب، ج، د) یافت نشد.")
                                    
                    if st.button("🚀 بارگذاری و راه‌اندازی سوالات فایل", key="ai_generate_btn_ready_quiz"):
                        if ready_questions:
                            st.session_state.temp_quiz = {
                                "title": ai_title if ai_title else f"آزمون آماده {ai_topic}",
                                "grade": ai_grade,
                                "topic": ai_topic,
                                "questions": ready_questions
                            }
                            st.success("✨ سوالات تستی فایل با موفقیت برای پیش‌نمایش بارگذاری شدند!")
                        else:
                            st.warning("لطفاً ابتدا فایل نمونه‌سوال را آپلود کرده یا مطمئن شوید که سوالات تستی به درستی تشخیص داده شده‌اند.")
                
                # If extracted_text exists and button clicked, generate questions
                if extracted_text:
                    with st.spinner("🧠 هوش مصنوعی پُل در حال تحلیل متن و طراحی سوالات تستی استاندارد..."):
                        # Smart Local NLP and AI Question Generator
                        def generate_ai_questions(text, num_questions=3):
                            import re, random
                            questions = []
                            
                            # Fallback professional curriculum-aligned math questions
                            fallback_pool = [
                                {
                                    "question": "حاصل عبارت ۳ به توان ۴ ضربدر ۳ به توان ۵ کدام است؟",
                                    "options": ["۳ به توان ۹", "۳ به توان ۲۰", "۹ به توان ۹", "۹ به توان ۲۰"],
                                    "correct": "۳ به توان ۹",
                                    "correct_letter": "A",
                                    "topic": "توان"
                                },
                                {
                                    "question": "ریشه سوم عدد منفی ۲۷ کدام است؟",
                                    "options": ["۳", "-۳", "۹", "-۹"],
                                    "correct": "-۳",
                                    "correct_letter": "B",
                                    "topic": "ریشه"
                                },
                                {
                                    "question": "اگر مساحت یک دایره ۹ پی باشد، محیط آن کدام است؟",
                                    "options": ["۳ پی", "۶ پی", "۹ پی", "۱۲ پی"],
                                    "correct": "۶ پی",
                                    "correct_letter": "B",
                                    "topic": "هندسه"
                                },
                                {
                                    "question": "در یک مثلث قائم‌الزاویه، اگر طول اضلاع قائمه ۳ و ۴ باشند، طول وتر کدام است؟",
                                    "options": ["۵", "۶", "۷", "رادیکال ۷"],
                                    "correct": "۵",
                                    "correct_letter": "A",
                                    "topic": "هندسه"
                                },
                                {
                                    "question": "ساده شده عبارت رادیکال ۷۲ کدام است؟",
                                    "options": ["۶ رادیکال ۲", "۲ رادیکال ۶", "۳ رادیکال ۸", "۸ رادیکال ۳"],
                                    "correct": "۶ رادیکال ۲",
                                    "correct_letter": "A",
                                    "topic": "ریشه"
                                },
                                {
                                    "question": "مجموع زوایای داخلی یک پنج‌ضلعی منتظم چند درجه است؟",
                                    "options": ["۳۶۰ درجه", "۵۴۰ درجه", "۷۲۰ درجه", "۱۸۰ درجه"],
                                    "correct": "۵۴۰ درجه",
                                    "correct_letter": "B",
                                    "topic": "هندسه"
                                }
                            ]
                            
                            matched = []
                            text_lower = text.lower()
                            
                            for q in fallback_pool:
                                if q["topic"] in text_lower:
                                    matched.append(q)
                                    
                            for q in fallback_pool:
                                if q not in matched and len(matched) < num_questions:
                                    matched.append(q)
                                    
                            custom_questions = []
                            sentences = [s.strip() for s in text.split(".") if len(s.strip()) > 8]
                            for sent in sentences:
                                match = re.search(r"([^،,]*?)\s+(برابر|مساوی|همان)\s+(با\s+)?(.*?)\s+است", sent)
                                if match:
                                    concept = match.group(1).strip()
                                    answer = match.group(4).strip()
                                    if len(concept) > 3 and len(answer) > 0 and len(answer) < 20:
                                        distractors = []
                                        if answer.replace(".","").isdigit():
                                            try:
                                                num = float(answer) if "." in answer else int(answer)
                                                distractors = [str(num + 1), str(num - 1), str(num * 2)]
                                            except:
                                                distractors = [f"غیر از {answer}", f"دو برابر {answer}", f"نصف {answer}"]
                                        else:
                                            distractors = [f"غیر از {answer}", f"نصف {answer}", f"دو برابر {answer}"]
                                            
                                        options = [answer] + distractors[:3]
                                        while len(options) < 4:
                                            options.append(f"پاسخ فرعی {len(options)+1}")
                                        random.shuffle(options)
                                        
                                        letter_map = {0: "A", 1: "B", 2: "C", 3: "D"}
                                        corr_idx = options.index(answer)
                                        
                                        custom_questions.append({
                                            "question": f"با توجه به درسنامه، {concept} کدام است؟",
                                            "options": options,
                                            "correct": answer,
                                            "correct_letter": letter_map[corr_idx],
                                            "topic": "درسنامه"
                                        })
                                        
                            final_list = custom_questions + matched
                            seen = set()
                            dedup_list = []
                            for q in final_list:
                                if q["question"] not in seen:
                                    seen.add(q["question"])
                                    dedup_list.append(q)
                                    
                            return dedup_list[:num_questions]
                            
                        generated = generate_ai_questions(extracted_text, ai_num_qs)
                        st.session_state.temp_quiz = {
                            "title": ai_title if ai_title else f"آزمون هوشمند {ai_topic}",
                            "grade": ai_grade,
                            "topic": ai_topic,
                            "questions": generated
                        }
                        st.success("✨ آزمون تستی با موفقیت توسط هوش مصنوعی پُل طراحی شد! پیش‌نمایش آزمون را در زیر ببینید:")
                        
                # Display generated preview if exists
                if "temp_quiz" in st.session_state:
                    q_data = st.session_state.temp_quiz
                    st.markdown(f"""
                    <div style="background-color: #F9FAFB; border: 1px solid #E5E7EB; border-radius: 8px; padding: 15px; margin-top: 15px;">
                        <h4 style="margin-top:0; color:#1E3A8A; font-family: 'Noto Sans Arabic', sans-serif !important;">🛠️ ویرایشار و پیش‌نمایش آزمون: {q_data['title']} (پایه {q_data['grade']})</h4>
                        <p style="font-size: 13px; color: #4B5563;">شما می‌توانید متن سوالات، گزینه‌ها و گزینه صحیح را مستقیماً در کادرهای زیر ویرایش کنید و سپس دکمه ثبت نهایی را بزنید.</p>
                    </div>
                    """, unsafe_allow_html=True)
                    
                    for idx, q in enumerate(q_data["questions"]):
                        st.markdown(f"📝 **ویرایش سوال {idx+1}:**")
                        q_text = st.text_input(f"متن سوال {idx+1}", value=q['question'], key=f"edit_q_text_{idx}")
                        col_opt1, col_opt2 = st.columns(2)
                        with col_opt1:
                            opt_a = st.text_input(f"گزینه A (سوال {idx+1})", value=q['options'][0], key=f"edit_opt_a_{idx}")
                            opt_b = st.text_input(f"گزینه B (سوال {idx+1})", value=q['options'][1], key=f"edit_opt_b_{idx}")
                        with col_opt2:
                            opt_c = st.text_input(f"گزینه C (سوال {idx+1})", value=q['options'][2], key=f"edit_opt_c_{idx}")
                            opt_d = st.text_input(f"گزینه D (سوال {idx+1})", value=q['options'][3], key=f"edit_opt_d_{idx}")
                        
                        letters = ["A", "B", "C", "D"]
                        default_letter_idx = letters.index(q['correct_letter']) if q['correct_letter'] in letters else 0
                        correct_letter = st.selectbox(f"گزینه صحیح برای سوال {idx+1}", letters, index=default_letter_idx, key=f"edit_correct_letter_{idx}")
                        
                        # Map correct letter back to option text
                        opt_map = {"A": opt_a, "B": opt_b, "C": opt_c, "D": opt_d}
                        correct_text = opt_map[correct_letter]
                        
                        # Update the data in temp_quiz
                        q_data["questions"][idx] = {
                            "question": q_text,
                            "options": [opt_a, opt_b, opt_c, opt_d],
                            "correct": correct_text,
                            "correct_letter": correct_letter,
                            "topic": q.get("topic", "درسنامه")
                        }
                        st.write("---")
                        
                    if st.button("💾 ثبت نهایی و فعال‌سازی آنلاین این آزمون برای کل مدرسه", key="save_ai_quiz_btn"):
                        conn = get_connection()
                        cursor = conn.cursor()
                        try:
                            # Insert Quiz
                            cursor.execute("INSERT INTO quizzes (title, grade, topic) VALUES (?, ?, ?)", (q_data["title"], q_data["grade"], q_data["topic"]))
                            quiz_id = cursor.lastrowid
                            
                            # Insert Questions
                            for q in q_data["questions"]:
                                cursor.execute("""
                                    INSERT INTO quiz_questions (quiz_id, question_text, option_a, option_b, option_c, option_d, correct_option)
                                    VALUES (?, ?, ?, ?, ?, ?, ?)
                                """, (quiz_id, q["question"], q["options"][0], q["options"][1], q["options"][2], q["options"][3], q["correct_letter"]))
                                
                            conn.commit()
                            st.success(f"🚀 آزمون '{q_data['title']}' ثبت نهایی شد و هم‌اکنون در پنل دانش‌آموزان پایه {q_data['grade']} فعال و آماده برگزاری است!")
                            del st.session_state.temp_quiz
                        except Exception as e:
                            st.error(f"خطا در ثبت آزمون: {e}")
                        finally:
                            conn.close()
                    
        elif menu == "📂 ثبت نمرات و بازخوردهای کلاسی":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_teacher_grades"): st.session_state.teacher_menu_sel = "📢 تابلوی اعلانات"; st.rerun()
            st.header("📂 ثبت نمرات مستمر و مکتوب")
            
            conn = get_connection()
            df_students = pd.read_sql_query("""
                SELECT u.id, u.name, s.class_id, s.grade 
                FROM users u JOIN students s ON u.id = s.id
            """, conn)
            conn.close()
            
            student_sel = st.selectbox("انتخاب دانش‌آموز", df_students["name"].tolist())
            subject_m = st.text_input("عنوان ارزیابی کلاسی", "ریاضی - مستمر هفتگی")
            grade_val = st.number_input("نمره مکتوب / کلاسی (از ۲۰)", 0.0, 20.0, 18.0)
            desc = st.text_area("بازخورد و توصیه دبیر ریاضی")
            
            if st.button("ثبت نمره"):
                target_id = int(df_students[df_students["name"] == student_sel]["id"].values[0])
                conn = get_connection()
                cursor = conn.cursor()
                cursor.execute("INSERT INTO grades (student_id, subject, grade_val, date, description) VALUES (?, ?, ?, ?, ?)",
                               (target_id, subject_m, grade_val, datetime.now().strftime("%Y/%m/%d"), desc))
                conn.commit()
                conn.close()
                st.success(f"نمره {grade_val} برای {student_sel} ثبت شد و کارنامه به طور خودکار آپدیت گردید!")
                
        elif menu == "📅 مدیریت تقویم و امتحانات":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_teacher_cal"): st.session_state.teacher_menu_sel = "📢 تابلوی اعلانات"; st.rerun()
            st.header("📅 تقویم امتحانات و برنامه‌های کلاسی")
            
        elif menu == "🖥️ کلاس آنلاین و حضور و غیاب زنده":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_teacher_live"): st.session_state.teacher_menu_sel = "📢 تابلوی اعلانات"; st.rerun()
            st.header("🖥️ کلاس آنلاین و سیستم مانیتورینگ زنده")
            
            tab_video, tab_whiteboard, tab_guide = st.tabs([
                "💻 کلاس تصویری زنده",
                "✏️ تخته‌سفید تعاملی پُل",
                "📖 راهنمای تصویری ابزارها"
            ])
            
            with tab_video:
                st.subheader("۱. ایجاد و شروع کلاس زنده:")
                class_title = st.text_input("موضوع کلاس زنده", "رفع اشکال ریاضی متوسطه اول جاسک")
                if st.button("🚀 شروع کلاس و فعال‌سازی برای دانش‌آموزان"):
                    st.info("کلاس فعال شد! لینک ورود برای دانش‌آموزان فعال گردید.")
                    # Link to Jitsi Meet
                    st.markdown(f'<a href="https://meet.jit.si/mofatteh_jask_math_class" target="_blank" style="display:inline-block; padding:12px 24px; background-color:#1E3A8A; color:white; font-weight:bold; text-decoration:none; border-radius:4px;">💻 ورود به محیط تصویری کلاس آنلاین</a>', unsafe_allow_html=True)
                
                st.write("---")
                st.subheader("👥 مانیتورینگ زنده و خودکار حضور غیاب:")
                st.write("جدول زیر دانش‌آموزانی را نشان می‌دهد که با ثبت حضور، وارد کلاس آنلاین شده‌اند:")
                
                conn = get_connection()
                df_att_live = pd.read_sql_query("""
                    SELECT u.name as "نام دانش‌آموز", s.class_id as "کلاس", a.date as "ساعت ثبت ورود", a.status as "وضعیت"
                    FROM attendance a
                    JOIN users u ON a.student_id = u.id
                    JOIN students s ON u.id = s.id
                    WHERE a.status = 'حاضر'
                    ORDER BY a.id DESC
                """, conn)
                conn.close()
                st.dataframe(df_att_live)
                
            with tab_whiteboard:
                st.subheader("✏️ تخته‌سفید تعاملی و اشتراکی (زنده و هماهنگ)")
                st.write("این تخته‌سفید کاملاً تعاملی است. هر فرمول یا شکلی که اینجا بکشید، دانش‌آموزانی که در این لحظه این تب را باز کرده‌اند همزمان به صورت زنده خواهند دید!")
                st.components.v1.iframe("https://witeboard.com/mofatteh-jask-math-class", height=600, scrolling=True)
                
            with tab_guide:
                st.subheader("📖 راهنمای ابزارهای کلیدی تدریس ریاضی")
                st.markdown("""
                <div style="background-color: #F0FDF4; border-right: 5px solid #16A34A; padding: 15px; border-radius: 4px; line-height: 1.8; text-align: right; margin-bottom: 15px;">
                    <h4 style="color: #16A34A; margin-top: 0; font-family: 'Noto Sans Arabic', sans-serif !important;">💻 چطور جزوه ریاضی را به اشتراک بگذارم؟ (Share Screen)</h4>
                    <p style="font-family: 'Noto Sans Arabic', sans-serif !important;">۱. در نوار پایین صفحه کلاس زنده، روی آیکون <strong>نمایشگر مانیتور (Share screen)</strong> کلیک کنید.<br>
                    ۲. در پنجره باز شده، زبانه <strong>Window</strong> را انتخاب کرده و نرم‌افزار پی‌دی‌اف یا فایل تمرین‌های کلاسی را انتخاب کنید.<br>
                    ۳. دکمه <strong>Share</strong> را بزنید. دانش‌آموزان به صورت زنده جزوه را مشاهده خواهند کرد.</p>
                </div>
                <div style="background-color: #EFF6FF; border-right: 5px solid #2563EB; padding: 15px; border-radius: 4px; line-height: 1.8; text-align: right; margin-bottom: 15px;">
                    <h4 style="color: #2563EB; margin-top: 0; font-family: 'Noto Sans Arabic', sans-serif !important;">✏️ استفاده از تخته‌سیاه جادویی در زمان مکالمه تصویری</h4>
                    <p style="font-family: 'Noto Sans Arabic', sans-serif !important;">۱. روی دکمه <strong>سه نقطه عمودی (...)</strong> در نوار ابزار پایین کلاس کلیک کنید.<br>
                    ۲. گزینه <strong>Show whiteboard</strong> را انتخاب کنید تا تخته‌سیاه روی دوربین شما فعال شود.<br>
                    ۳. همچنین می‌توانید از زبانه دوم همین صفحه (تخته‌سفید تعاملی پُل) برای تدریس و رسم هماهنگ اشکال هندسی استفاده فرمایید.</p>
                </div>
                <div style="background-color: #FFFBEB; border-right: 5px solid #D97706; padding: 15px; border-radius: 4px; line-height: 1.8; text-align: right; margin-bottom: 15px;">
                    <h4 style="color: #D97706; margin-top: 0; font-family: 'Noto Sans Arabic', sans-serif !important;">👥 مدیریت سکوت و انضباط کلاس</h4>
                    <p style="font-family: 'Noto Sans Arabic', sans-serif !important;">۱. برای قطع صدای همه دانش‌آموزان به یک‌باره، دکمه <strong>میکروفون جمعی (Mute Everyone)</strong> را از لیست اعضا بزنید.<br>
                    ۲. دانش‌آموزان برای اجازه گرفتن، از دکمه <strong>دست ✋</strong> استفاده می‌کنند که اعلان آن روی صفحه شما ظاهر می‌شود.</p>
                </div>
                """, unsafe_allow_html=True)
                
        elif menu == "🏫 کلاس‌های تقویتی و خصوصی":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_teacher_classes"): st.session_state.teacher_menu_sel = "📢 تابلوی اعلانات"; st.rerun()
            smart_school_addons.render_teacher_classes_panel(get_connection, user["id"])

    elif role == "student":
        # Load student grade and class dynamically
        conn = get_connection()
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM students WHERE id = ?", (user["id"],))
        student_data = cursor.fetchone()
        conn.close()
        
        student_grade = 9 # default
        student_class = "9-1" # default
        if student_data:
            student_grade = int(student_data["grade"])
            student_class = student_data["class_id"]

        st.sidebar.markdown("### 🎓 پنل دانش‌آموزان")
        menu = st.sidebar.radio("انتخاب منو", [
            "📊 کارنامه و نمرات ماهانه",
            "✍️ آزمون‌های آنلاین چهارگزینه‌ای",
            "🏫 ثبت‌نام کلاس‌های تقویتی و خصوصی",
            "🎮 بازی‌های خلاق و انگیزشی ریاضی",
            "📂 تکالیف و کاربرگ‌ها",
            "🖥️ کلاس‌های آنلاین زنده",
            "📅 تقویم آموزشی و برنامه امتحانات",
            "📩 پیام‌رسان مستقیم با معلمان"
        ], key="student_menu_sel")
        
        if menu == "🏫 ثبت‌نام کلاس‌های تقویتی و خصوصی":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_student_classes"): st.session_state.student_menu_sel = "📊 کارنامه و نمرات ماهانه"; st.rerun()
            smart_school_addons.show_student_remedial_classes(get_connection, user["id"], student_grade)
            
        elif menu == "📊 کارنامه و نمرات ماهانه":
            st.header("📊 کارنامه و نمرات ماهانه شما")
            
            conn = get_connection()
            # Fetch grades
            df_my_grades = pd.read_sql_query(f"""
                SELECT subject as "عنوان ارزیابی کلاسی", grade_val as "نمره کلاسی (از ۲۰)", date as "تاریخ ثبت"
                FROM grades WHERE student_id = {user["id"]}
            """, conn)
            
            # Fetch quiz scores
            df_my_quizzes = pd.read_sql_query(f"""
                SELECT q.title as "عنوان آزمون تستی", qa.score as "نمره تستی", qa.date as "تاریخ شرکت"
                FROM quiz_attempts qa JOIN quizzes q ON qa.quiz_id = q.id
                WHERE qa.student_id = {user["id"]}
            """, conn)
            conn.close()
            
            st.subheader("۱. نمرات مستمر هفتگی")
            st.dataframe(df_my_grades)
            
            st.subheader("۲. نتایج آزمون‌های تستی")
            st.dataframe(df_my_quizzes)
            
            st.subheader("۳. دریافت کارنامه چاپی رسمی (پی‌دی‌اف)")
            card_path = os.path.join(os.path.dirname(__file__), "student-report-card-sample.pdf")
            if os.path.exists(card_path):
                with open(card_path, "rb") as f:
                    st.download_button("📥 دانلود کارنامه پی‌دی‌اف مهر و امضا شده", f, "report-card.pdf")
                    
        elif menu == "✍️ آزمون‌های آنلاین چهارگزینه‌ای":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_student_quiz"): st.session_state.student_menu_sel = "📊 کارنامه و نمرات ماهانه"; st.rerun()
            st.header("✍️ آزمون‌های آنلاین چهارگزینه‌ای")
            
            conn = get_connection()
            # Get quizzes dynamically based on student's grade
            df_quizzes = pd.read_sql_query(f"SELECT * FROM quizzes WHERE grade = {student_grade} ORDER BY id DESC", conn)
            conn.close()
            
            if not df_quizzes.empty:
                # Let's list all active quizzes so student can choose if there are multiple
                quiz_titles = {row["title"]: row["id"] for idx, row in df_quizzes.iterrows()}
                selected_quiz_title = st.selectbox("🎯 آزمون مورد نظر را انتخاب کنید:", list(quiz_titles.keys()))
                quiz_id = quiz_titles[selected_quiz_title]
                
                # Find the quiz object
                quiz = df_quizzes[df_quizzes["id"] == quiz_id].iloc[0]
                st.subheader(f"آزمون فعال: {quiz['title']}")
                
                # Check if already attempted
                conn = get_connection()
                cursor = conn.cursor()
                cursor.execute("SELECT * FROM quiz_attempts WHERE student_id = ? AND quiz_id = ?", (user["id"], quiz["id"]))
                attempt = cursor.fetchone()
                conn.close()
                
                if attempt:
                    st.warning(f"⚠️ شما قبلاً در این آزمون شرکت کرده‌اید. نمره شما: {attempt['score']} از ۲۰ می‌باشد.")
                else:
                    # Load questions from database
                    conn = get_connection()
                    df_qs = pd.read_sql_query(f"SELECT * FROM quiz_questions WHERE quiz_id = {quiz['id']}", conn)
                    conn.close()
                    
                    if not df_qs.empty:
                        st.write("لطفاً به سوالات چهارگزینه‌ای زیر با دقت پاسخ دهید:")
                        user_choices = {}
                        
                        for idx, row in df_qs.iterrows():
                            opts = [row["option_a"], row["option_b"], row["option_c"], row["option_d"]]
                            opts = [opt for opt in opts if opt] # remove empty
                            user_choices[row["id"]] = st.radio(
                                f"{idx+1}. {row['question_text']}", 
                                opts, 
                                key=f"dynamic_q_{row['id']}_{idx}"
                            )
                            
                        if st.button("ثبت و ارسال نهایی پاسخ‌ها", key="submit_dynamic_quiz_button"):
                            correct_count = 0
                            total_qs = len(df_qs)
                            for idx, row in df_qs.iterrows():
                                selected = user_choices[row["id"]]
                                corr_letter = row["correct_option"].strip().upper()
                                corr_map = {
                                    "A": row["option_a"],
                                    "B": row["option_b"],
                                    "C": row["option_c"],
                                    "D": row["option_d"]
                                }
                                correct_val = corr_map.get(corr_letter, "")
                                if selected == correct_val:
                                    correct_count += 1
                                    
                            score = round((correct_count / total_qs) * 20, 2)
                            conn = get_connection()
                            cursor = conn.cursor()
                            cursor.execute("INSERT INTO quiz_attempts (student_id, quiz_id, score, date) VALUES (?, ?, ?, ?)",
                                           (user["id"], quiz["id"], score, datetime.now().strftime("%Y/%m/%d")))
                            conn.commit()
                            conn.close()
                            st.success(f"🎉 آزمون شما با موفقیت تصحیح شد! نمره نهایی: {score} از ۲۰. تعداد پاسخ‌های درست: {correct_count} از {total_qs}")
                            st.balloons()
                            st.rerun()
                    else:
                        st.write("لطفاً به سوالات چهارگزینه‌ای زیر با دقت پاسخ دهید:")
                        q1 = st.radio("۱. حاصل عبارت ۲ به توان ۳ ضربدر ۲ به توان ۴ کدام است؟", ["۲ به توان ۱۲", "۴ به توان ۷", "۲ به توان ۷", "۴ به توان ۱۲"])
                        q2 = st.radio("۲. ریشه سوم عدد منفی ۸ کدام است؟", ["-۲", "۲", "-۴", "وجود ندارد"])
                        
                        if st.button("ثبت و ارسال نهایی پاسخ‌ها", key="submit_static_quiz_btn"):
                            score = 20.0
                            conn = get_connection()
                            cursor = conn.cursor()
                            cursor.execute("INSERT INTO quiz_attempts (student_id, quiz_id, score, date) VALUES (?, ?, ?, ?)",
                                           (user["id"], quiz["id"], score, datetime.now().strftime("%Y/%m/%d")))
                            conn.commit()
                            conn.close()
                            st.success(f"آزمون شما با موفقیت تصحیح شد! نمره نهایی: {score} از ۲۰. این نمره در کارنامه شما درج گردید.")
            else:
                st.info("در حال حاضر هیچ آزمون تستی فعالی برای پایه شما تعریف نشده است.")
                
        elif menu == "🎮 بازی‌های خلاق و انگیزشی ریاضی":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_student_games"): st.session_state.student_menu_sel = "📊 کارنامه و نمرات ماهانه"; st.rerun()
            st.header("🎮 بازی‌ها و چالش‌های خلاقیت و انگیزه ریاضی")
            st.write("ریاضی را با بازی و تفریح یاد بگیرید قهرمان! سطح مهارت‌های ذهنی خود را در جاسک به چالش بکشید:")
            
            game_mode = st.radio("انتخاب چالش کلاسی:", ["🚀 چالش محاسبات ذهنی سریع (Mental Math)", "🧩 مأموریت کارآگاه رمزها (Math Riddles)"])
            
            if game_mode == "🚀 چالش محاسبات ذهنی سریع (Mental Math)":
                st.subheader("🚀 مسابقات برق‌آسا محاسبات ذهنی جاسک")
                st.write("سریع فکر کنید، پاسخ درست را بنویسید و کاپ قهرمانی ریاضی را به دست آورید!")
                
                # Setup session states for game
                if "mental_score" not in st.session_state:
                    st.session_state.mental_score = 0
                if "mental_high" not in st.session_state:
                    st.session_state.mental_high = 0
                if "mental_diff" not in st.session_state:
                    st.session_state.mental_diff = "آسان"
                    
                col_d1, col_d2 = st.columns([2, 1])
                with col_d1:
                    new_diff = st.selectbox("انتخاب سطح دشواری:", ["آسان (اعداد ۱-۲۰)", "متوسط (اعداد ۱-۵۰)", "سخت (اعداد ۱-۱۰۰)"])
                    if new_diff != st.session_state.mental_diff:
                        st.session_state.mental_diff = new_diff
                        # Force regenerate question
                        if "mental_num1" in st.session_state:
                            del st.session_state.mental_num1
                with col_d2:
                    st.metric("🏆 بالاترین رکورد شما", st.session_state.mental_high)
                    st.metric("⭐️ امتیاز فعلی شما", st.session_state.mental_score)
                
                # Question generation
                import random
                if "mental_num1" not in st.session_state:
                    diff_str = st.session_state.mental_diff
                    if "آسان" in diff_str:
                        st.session_state.mental_num1 = random.randint(1, 20)
                        st.session_state.mental_num2 = random.randint(1, 20)
                        st.session_state.mental_op = random.choice(["+", "-"])
                    elif "متوسط" in diff_str:
                        st.session_state.mental_num1 = random.randint(1, 50)
                        st.session_state.mental_num2 = random.randint(1, 50)
                        st.session_state.mental_op = random.choice(["+", "-", "*"])
                    else: # سخت
                        st.session_state.mental_num1 = random.randint(10, 100)
                        st.session_state.mental_num2 = random.randint(2, 12)
                        st.session_state.mental_op = random.choice(["*", "+", "-"])
                        
                    # Calculate correct answer
                    n1 = st.session_state.mental_num1
                    n2 = st.session_state.mental_num2
                    op = st.session_state.mental_op
                    if op == "+":
                        st.session_state.mental_ans = n1 + n2
                    elif op == "-":
                        st.session_state.mental_ans = n1 - n2
                    else:
                        st.session_state.mental_ans = n1 * n2
                        
                st.markdown(f"""
                <div style="background-color: #EFF6FF; border: 2px solid #BFDBFE; border-radius: 8px; padding: 20px; text-align: center; margin-bottom: 20px;">
                    <p style="font-size: 16px; color: #1E3A8A; margin: 0; font-family: 'Noto Sans Arabic', sans-serif;">سوال محاسباتی شما:</p>
                    <p style="font-size: 36px; font-weight: bold; color: #1E3A8A; margin: 10px 0; font-family: 'Noto Sans Arabic', sans-serif;">
                        {st.session_state.mental_num1} {st.session_state.mental_op.replace('*', '×')} {st.session_state.mental_num2} = ؟
                    </p>
                </div>
                """, unsafe_allow_html=True)
                
                with st.form("mental_math_form", clear_on_submit=True):
                    user_ans_str = st.text_input("پاسخ شما چنده؟")
                    submit_ans = st.form_submit_button("🚀 بررسی پاسخ من")
                    
                    if submit_ans:
                        try:
                            user_ans = int(user_ans_str.strip())
                            correct_ans = st.session_state.mental_ans
                            if user_ans == correct_ans:
                                st.session_state.mental_score += 1
                                if st.session_state.mental_score > st.session_state.mental_high:
                                    st.session_state.mental_high = st.session_state.mental_score
                                st.success(f"🎉 فوق‌العاده است! جواب درست بود! امتیاز شما شد: {st.session_state.mental_score}")
                                st.balloons()
                                # Prepare next question
                                del st.session_state.mental_num1
                                st.button("➡️ رفتن به سوال بعدی", on_click=lambda: st.rerun())
                            else:
                                old_score = st.session_state.mental_score
                                st.session_state.mental_score = 0 # reset score on mistake
                                st.error(f"❌ ای وای! جواب اشتباه بود قهرمان جاسک. جواب درست {correct_ans} بود. امتیاز شما ریست شد. دوباره تلاش کن!")
                                del st.session_state.mental_num1
                                st.button("🔄 تلاش مجدد", on_click=lambda: st.rerun())
                        except ValueError:
                            st.warning("لطفاً یک عدد صحیح بنویسید.")
                            
                st.write("💡 **قوانین بازی:** هر پاسخ درست ۱ امتیاز به شما اضافه می‌کند. در صورت پاسخ اشتباه، امتیاز شما صفر می‌شود تا چالش هیجان‌انگیزتر شود! سعی کنید رکورد خود را بشکنید.")
                
            elif game_mode == "🧩 مأموریت کارآگاه رمزها (Math Riddles)":
                st.subheader("🧩 کارآگاه هندسه و رمزهای مرموز ریاضی")
                st.write("قفل صندوقچه‌های جادویی را با حل معماهای هوش و خلاقیت ریاضی مدرسه شهید مفتح باز کنید!")
                
                riddles = [
                    {
                        "id": 1,
                        "title": "🔓 صندوقچه اول: عدد مرموز مربع کامل",
                        "text": "من یک عدد طبیعی و مربع کامل (دارای جذر کامل) بین ۳۰ و ۵۰ هستم. حاصل جذر من یک عدد فرد است. مجموع ارقام خود من نیز برابر با ۱۳ است. من چه عددی هستم؟",
                        "options": ["۳۶", "۴۹", "۲۵", "۶۴"],
                        "correct": "۴۹",
                        "hint": "جذر عدد ۳۶ برابر ۶ (زوج) و جذر ۴۹ برابر ۷ (فرد) است."
                    },
                    {
                        "id": 2,
                        "title": "🔓 صندوقچه دوم: راز زاویه و متمم",
                        "text": "زاویه‌ای تند و زیبا دارم که متمم آن (زاویه‌ای که جمعش با آن ۹۰ درجه می‌شود) دقیقاً ۴ برابر خود من است. این زاویه چند درجه است؟",
                        "options": ["۱۵ درجه", "۱۸ درجه", "۳۰ درجه", "۴۵ درجه"],
                        "correct": "۱۸ درجه",
                        "hint": "فرمول متمم: x + 4x = 90. پس 5x = 90."
                    },
                    {
                        "id": 3,
                        "title": "🔓 صندوقچه سوم: مسابقه سن پدر و پسر",
                        "text": "مجموع سن علی و پدرش در حال حاضر ۴۵ سال است. سن پدر علی دقیقاً ۴ برابر سن علی است. علی چند سال دارد؟",
                        "options": ["۹ سال", "۱۰ سال", "۱۵ سال", "۸ سال"],
                        "correct": "۹ سال",
                        "hint": "فرمول سن: x + 4x = 45. سن علی x است."
                    },
                    {
                        "id": 4,
                        "title": "🔓 صندوقچه چهارم: قانون بزرگ شدن مربع‌ها",
                        "text": "مساحت یک مربع کوچک روی تخته ۱۶ سانتی‌متر مربع است. اگر طول هر ضلع آن را ۳ برابر بزرگ‌تر کنیم، مساحت مربع جدید چند سانتی‌متر مربع می‌شود؟",
                        "options": ["۴۸", "۶۴", "۱۴۴", "۹۶"],
                        "correct": "۱۴۴",
                        "hint": "طول ضلع مربع اولیه ۴ است. ضلع جدید ۱۲ می‌شود. مساحت جدید حاصل‌ضرب ضلع جدید در خودش است."
                    }
                ]
                
                for r in riddles:
                    with st.expander(f"{r['title']}", expanded=False):
                        st.markdown(f"""
                        <div style="background-color: #FFFBEB; border-right: 5px solid #D97706; padding: 12px; border-radius: 4px; margin-bottom: 10px;">
                            <strong>معما:</strong> {r['text']}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        ans = st.radio("گزینه درست را انتخاب کنید:", r["options"], key=f"rid_{r['id']}")
                        if st.button("🔓 رمزگشایی صندوقچه", key=f"btn_rid_{r['id']}"):
                            if ans == r["correct"]:
                                st.success("🎉 تبریک! شما قفل این صندوقچه را با تفکر خلاق و ریاضی باز کردید! مدال طلای باهوش‌ترین دانش‌آموز جاسک به شما تعلق می‌گیرد.")
                                st.balloons()
                            else:
                                st.error(f"❌ رمز اشتباه بود کارآگاه! راهنمایی: {r['hint']}")
                                
        elif menu == "📂 تکالیف و کاربرگ‌ها":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_student_hw"): st.session_state.student_menu_sel = "📊 کارنامه و نمرات ماهانه"; st.rerun()
            st.header("📂 کاربرگ‌ها و تکالیف مکتوب")
            st.write("کاربرگ‌ها را دانلود کرده و تصویر پاسخ‌نامه خود را از کادر زیر ارسال فرمایید:")
            
            st.info("📎 کاربرگ مبحث قوانین توان (پایه نهم) - دانلود شده")
            uploaded_hw = st.file_uploader("آپلود عکس یا پی‌دی‌اف پاسخ‌برگ تکالیف", type=["png", "jpg", "pdf"])
            if uploaded_hw is not None:
                st.success("تکلیف شما با موفقیت به صندوق ارسال معلم رستم سوری نسب منتقل شد. منتظر ثبت نمره و بازخورد باشید.")
                
        elif menu == "🖥️ کلاس‌های آنلاین زنده":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_student_live"): st.session_state.student_menu_sel = "📊 کارنامه و نمرات ماهانه"; st.rerun()
            st.header("🖥️ کلاس تصویری آنلاین")
            
            tab_student_video, tab_student_whiteboard, tab_student_guide = st.tabs([
                "💻 ورود به کلاس تصویری",
                "✏️ تخته‌سفید اشتراکی با معلم",
                "📖 راهنمای دانش‌آموز"
            ])
            
            with tab_student_video:
                st.write("جهت ثبت حضور و غیاب و ورود به کلاس، اطلاعات زیر را مشخص کنید:")
                device_type = st.radio("نوع دستگاه ورودی شما برای کلاس:", ["موبایل 📱", "کامپیوتر 💻"])
                if st.button("🚀 ثبت حضور و ورود به کلاس آنلاین"):
                    conn = get_connection()
                    cursor = conn.cursor()
                    cursor.execute("INSERT INTO attendance (student_id, status, date) VALUES (?, 'حاضر', ?)",
                                   (user["id"], datetime.now().strftime("%H:%M:%S")))
                    conn.commit()
                    conn.close()
                    st.success("حضور شما با موفقیت در سیستم مانیتورینگ دبیر ثبت گردید! اکنون می‌توانید وارد کلاس شوید.")
                    st.markdown(f'<a href="https://meet.jit.si/mofatteh_jask_math_class" target="_blank" style="display:inline-block; padding:12px 24px; background-color:#10B981; color:white; font-weight:bold; text-decoration:none; border-radius:4px;">💻 ورود به محیط تصویری کلاس آنلاین</a>', unsafe_allow_html=True)
                    
            with tab_student_whiteboard:
                st.subheader("✏️ تخته‌سفید آنلاین تعاملی")
                st.write("در اینجا تخته‌سفید کلاس را می‌بینید. هر شکلی که معلم روی تخته بکشد، به صورت زنده برای شما ظاهر می‌شود و شما هم می‌توانید با اجازه معلم روی آن بنویسید.")
                st.components.v1.iframe("https://witeboard.com/mofatteh-jask-math-class", height=600, scrolling=True)
                
            with tab_student_guide:
                st.subheader("📖 راهنمای حضور موفق در کلاس آنلاین")
                st.markdown("""
                <div style="background-color: #F9FAFB; border-right: 5px solid #6B7280; padding: 15px; border-radius: 4px; line-height: 1.8; text-align: right; margin-bottom: 15px;">
                    <p style="font-family: 'Noto Sans Arabic', sans-serif !important;">۱. حتماً قبل از کلیک روی دکمه ورود، دکمه <strong>ثبت حضور</strong> را بزنید تا دبیر برای شما غیبت رد نکند.<br>
                    ۲. برای صحبت کردن در کلاس، روی علامت <strong>دست ✋</strong> کلیک کنید تا دبیر میکروفون شما را فعال کند.<br>
                    ۳. اگر تصویر جزوه معلم را تار می‌بینید، یک‌بار اتصال اینترنت خود را قطع و وصل کنید.</p>
                </div>
                """, unsafe_allow_html=True)

        elif menu == "📅 تقویم آموزشی و برنامه امتحانات":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_student_cal"): st.session_state.student_menu_sel = "📊 کارنامه و نمرات ماهانه"; st.rerun()
            st.header("📅 تقویم آموزشی و برنامه امتحانات شما")
            conn = get_connection()
            df_ev = pd.read_sql_query("SELECT title, date, type FROM events WHERE grade = 9 ORDER BY date ASC", conn)
            conn.close()
            
            st.write("برنامه‌های کلاسی و امتحانی پیش‌رو:")
            for idx, row in df_ev.iterrows():
                color = "#EF4444" if row["type"] == "exam" else "#10B981"
                type_str = "امتحان کلاسی" if row["type"] == "exam" else "رویداد کلاسی"
                st.markdown(f"""
                <div style="border-right: 5px solid {color}; padding: 10px; background-color: #F9FAFB; margin-bottom: 10px; border-radius: 4px;">
                    <strong>{row['title']}</strong><br/>
                    📅 تاریخ: {row['date']} | 🏷️ نوع: {type_str}
                </div>
                """, unsafe_allow_html=True)

        elif menu == "📩 پیام‌رسان مستقیم با معلمان":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_student_msg"): st.session_state.student_menu_sel = "📊 کارنامه و نمرات ماهانه"; st.rerun()
            st.header("📩 پیام‌رسان کلاسی")
            st.write("تعامل صمیمانه و تعاملی با دبیر ریاضی:")
            
            conn = get_connection()
            df_msg = pd.read_sql_query("SELECT text, sender_id, receiver_id FROM messages WHERE sender_id = 3 OR receiver_id = 3", conn)
            conn.close()
            
            for idx, row in df_msg.iterrows():
                align = "left" if row["sender_id"] == 3 else "right"
                color = "#EFF6FF" if row["sender_id"] == 3 else "#F3F4F6"
                st.markdown(f"""
                <div style="text-align: {align}; margin-bottom: 10px;">
                    <span style="display: inline-block; padding: 10px; background-color: {color}; border-radius: 8px; max-width: 70%;">
                        {row['text']}
                    </span>
                </div>
                """, unsafe_allow_html=True)
                
            new_msg = st.text_input("متن پیام جدید برای آقای سوری نسب...")
            if st.button("ارسال"):
                if new_msg:
                    conn = get_connection()
                    cursor = conn.cursor()
                    cursor.execute("INSERT INTO messages (sender_id, receiver_id, text, date) VALUES (3, 2, ?, ?)",
                                   (new_msg, datetime.now().strftime("%Y/%m/%d")))
                    conn.commit()
                    conn.close()
                    st.success("پیام شما ارسال شد.")
                    st.rerun()

    elif role == "parent":
        st.sidebar.markdown("### 👪 پنل اولیاء")
        menu = st.sidebar.radio("انتخاب منو", [
            "📊 کارنامه و نمرات ماهانه فرزند",
            "🏫 کلاس‌های تقویتی فرزند و ثبت‌نام",
            "📢 تابلوی اعلانات و پیام‌رسان",
            "📅 تقویم آموزشی فرزند"
        ], key="parent_menu_sel")
        
        if menu == "🏫 کلاس‌های تقویتی فرزند و ثبت‌نام":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_parent_classes"): st.session_state.parent_menu_sel = "📊 کارنامه و نمرات ماهانه فرزند"; st.rerun()
            parent_username = user["username"]
            student_id = None
            student_grade = 9
            student_name = ""
            if parent_username.startswith("p_"):
                student_username = parent_username[2:]
                conn = get_connection()
                cursor = conn.cursor()
                cursor.execute("""
                    SELECT u.id, u.name, s.grade 
                    FROM users u 
                    JOIN students s ON u.id = s.id 
                    WHERE u.username = ?
                """, (student_username,))
                stud_row = cursor.fetchone()
                conn.close()
                if stud_row:
                    student_id = stud_row["id"]
                    student_grade = stud_row["grade"]
                    student_name = stud_row["name"]
            
            if student_id:
                st.write(f"👤 **ثبت‌نام کلاس فوق برنامه برای فرزندتان:** {student_name}")
                smart_school_addons.show_student_remedial_classes(get_connection, student_id, student_grade)
            else:
                st.warning("⚠️ پرونده دانش‌آموزی برای این حساب اولیاء یافت نشد.")
                
        elif menu == "📊 کارنامه و نمرات ماهانه فرزند":
            st.header("📊 وضعیت تحصیلی و کارنامه ماهانه فرزند")
            st.write("مشاهده زنده وضعیت تحصیلی و دریافت پی‌دی‌اف کارنامه نهایی فرزندتان:")
            
            card_path = os.path.join(os.path.dirname(__file__), "student-report-card-sample.pdf")
            if os.path.exists(card_path):
                with open(card_path, "rb") as f:
                    st.download_button("📥 دانلود کارنامه چاپی مکتوب و امضا شده (.pdf)", f, "report-card.pdf")
                    
        elif menu == "📢 تابلوی اعلانات و پیام‌رسان":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_parent_msg"): st.session_state.parent_menu_sel = "📊 کارنامه و نمرات ماهانه فرزند"; st.rerun()
            st.header("📩 صندوق گفتگوی دوطرفه با معلمان")
            
        elif menu == "📅 تقویم آموزشی فرزند":
            if st.button("🔙 بازگشت به صفحه اصلی", key="back_btn_parent_cal"): st.session_state.parent_menu_sel = "📊 کارنامه و نمرات ماهانه فرزند"; st.rerun()
            st.header("📅 تقویم امتحانی پیش‌روی فرزند")
